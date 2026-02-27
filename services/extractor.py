"""
문서 텍스트 추출 모듈
- PDF 파일에서 텍스트 추출 (PyPDF2)
- DOCX 파일에서 텍스트 추출 (python-docx)
"""

from PyPDF2 import PdfReader
from docx import Document


def extract_text_from_pdf(file_path: str) -> str:
    """
    PDF 파일에서 텍스트를 추출한다.

    Args:
        file_path: PDF 파일 경로

    Returns:
        추출된 텍스트 문자열
    """
    reader = PdfReader(file_path)
    text_parts = []

    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text_parts.append(page_text)

    return "\n".join(text_parts)


def extract_text_from_docx(file_path: str) -> str:
    """
    DOCX 파일에서 텍스트를 추출한다.

    Args:
        file_path: DOCX 파일 경로

    Returns:
        추출된 텍스트 문자열
    """
    doc = Document(file_path)
    text_parts = []

    # 본문 단락 추출
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            text_parts.append(paragraph.text.strip())

    # 표(Table) 내부 텍스트도 추출
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text_parts.append(cell.text.strip())

    return "\n".join(text_parts)


def extract_text(file_path: str, filename: str) -> str:
    """
    파일 확장자에 따라 적절한 추출기를 호출한다.

    Args:
        file_path: 파일 경로
        filename: 원본 파일명 (확장자 판별용)

    Returns:
        추출된 텍스트 문자열

    Raises:
        ValueError: 지원하지 않는 파일 형식일 때
    """
    lower_name = filename.lower()

    if lower_name.endswith(".pdf"):
        return extract_text_from_pdf(file_path)
    elif lower_name.endswith(".docx"):
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"지원하지 않는 파일 형식입니다: {filename}")
